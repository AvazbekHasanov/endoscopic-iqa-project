"""
Generate optimized statistical report for image quality metrics from the database.

Creates a comprehensive Word document with:
1. Executive Summary Dashboard
   - Dataset overview
   - Quality distribution (Good/Fair/Poor)
   - Key findings

2. Correlation Analysis
   - Heatmap showing relationships between metrics
   - Identifies redundant or complementary metrics

3. Detailed Metric Analysis
   - 5 key statistics per metric (reduced from 9 for clarity):
     * Count: Sample size
     * Mean: Average value
     * Std Dev: Variability/consistency
     * Min: Lower bound
     * Max: Upper bound
   - Distribution histogram
   - Boxplot for outlier detection

Best practices implemented:
- Dashboard-ready metrics
- Focus on actionable insights
- Reduced cognitive load (5 vs 9 statistics)
- Visual correlation analysis
- Quality classification for decision making
"""

import os
import sys
import psycopg2
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
from docx import Document
from docx.shared import Inches

# Add parent directory to path for imports
sys.path.append(str(Path(__file__).parent.parent))

from scripts.db_config import DB_CONFIG


class StatisticsReportGenerator:
    """Generate comprehensive statistical report for image quality metrics"""
    
    def __init__(self, db_config, output_dir='reports', n_sample_images=3, histogram_bins='auto'):
        """
        Initialize the StatisticsReportGenerator.
        
        Args:
            db_config (dict): Database configuration with host, port, database, user, password
            output_dir (str): Directory to save the report and visualizations (default: 'reports')
            n_sample_images (int): Number of sample images to include in report (default: 3)
            histogram_bins (int or str): Number of bins for histograms or 'auto' for automatic (default: 'auto')
        """
        self.db_config = db_config
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.conn = None
        self.df = None
        self.n_sample_images = n_sample_images
        self.histogram_bins = histogram_bins  # Can be int or 'auto'
        
        # Metric names mapping (database column names)
        self.metrics = {
            'laplacian_variance': 'laplacian',
            'gradient_energy': 'gradient_energy',
            'rms_contrast': 'rms_contrast',
            'entropy': 'entropy',
            'noise_estimate': 'noise',
            'tenengrad': 'tenengrad',
            'mscn_std': 'mscn'
        }
        
        # Metric descriptions in Uzbek
        self.metric_descriptions = {
            'laplacian': "Laplacian o'lchovi tasvirning burchak va kontur aniqligini ko'rsatadi. Yuqori qiymat tasvir aniqroq ekanini bildiradi.",
            'gradient_energy': "Gradient energy tasvirning tekstura va qirralarining aniqligini baholaydi. Yuqori qiymatlar ko'proq qirralar va teksturani bildiradi.",
            'rms_contrast': "RMS kontrast tasvir yorqinligi va kontrast darajasini o'lchaydi. Yuqori RMS kontrast tasvir ko'proq kontrastli ekanini bildiradi.",
            'entropy': "Entropiya tasvirdagi axborot miqdorini ifodalaydi. Yuqori entropiya ko'proq tafsilot va murakkablikni bildiradi.",
            'noise': "Noise tasvirdagi shovqin darajasini o'lchaydi. Yuqori qiymat tasvirda ko'p shovqin borligini bildiradi.",
            'tenengrad': "Tenengrad o'lchovi tasvirning o'rtacha aniqligini ko'rsatadi. Yuqori qiymat tasvir aniqroq ekanini bildiradi.",
            'mscn': "MSCN tasvirning lokal kontrastini baholaydi. Yuqori qiymat ko'proq lokal kontrastni bildiradi."
        }
        
    def connect_database(self):
        """Connect to PostgreSQL database"""
        try:
            self.conn = psycopg2.connect(
                host=self.db_config.get('host', 'localhost'),
                port=self.db_config.get('port', 5432),
                database=self.db_config.get('database', 'postgres'),
                user=self.db_config.get('user', 'postgres'),
                password=self.db_config.get('password', '')
            )
            print(f"✓ Connected to database: {self.db_config.get('database')}")
        except Exception as e:
            print(f"✗ Error connecting to database: {str(e)}")
            sys.exit(1)
    
    def fetch_data(self):
        """Fetch image quality metrics from database"""
        query = """
        SELECT 
            id,
            laplacian_variance,
            gradient_energy,
            rms_contrast,
            entropy,
            noise_estimate,
            tenengrad,
            mscn_std,
            file_path
        FROM image_quality_metrics_hybrid
        """
        
        try:
            self.df = pd.read_sql(query, self.conn)
            print(f"✓ Fetched {len(self.df)} records from database")
        except Exception as e:
            print(f"✗ Error fetching data: {str(e)}")
            sys.exit(1)
        finally:
            if self.conn:
                self.conn.close()
    
    def calculate_statistics(self):
        """
        Calculate the 5 most useful statistics for each metric.
        
        Based on best practices for dashboard metrics:
        - count: Sample size (data quality indicator)
        - mean: Central tendency (average quality)
        - std: Variability (consistency indicator)
        - min: Lower bound (worst case)
        - max: Upper bound (best case)
        """
        stats_summary = {}
        
        for db_col, metric_name in self.metrics.items():
            if db_col in self.df.columns:
                data = self.df[db_col].dropna()
                
                if len(data) > 0:
                    stats_summary[metric_name] = {
                        'count': int(data.count()),
                        'mean': float(data.mean()),
                        'std': float(data.std()),
                        'min': float(data.min()),
                        'max': float(data.max())
                    }
                else:
                    print(f"⚠ No data available for {metric_name}")
        
        return stats_summary
    
    def create_visualizations(self, metric_name, db_col):
        """Create histogram and boxplot for a metric"""
        if db_col not in self.df.columns:
            print(f"⚠ Column {db_col} not found in dataframe")
            return None, None
        
        data = self.df[db_col].dropna()
        
        if len(data) == 0:
            print(f"⚠ No data available for {metric_name}")
            return None, None
        
        # Determine bins (use auto or configured value)
        bins = self.histogram_bins if self.histogram_bins != 'auto' else min(50, max(10, len(data) // 100))
        
        # Create histogram
        plt.figure(figsize=(6, 4))
        sns.histplot(data, bins=bins, kde=True, color='skyblue')
        plt.title(f"{metric_name} distribution")
        plt.xlabel(metric_name)
        plt.ylabel("Frequency")
        hist_path = self.output_dir / f"{metric_name}_hist.png"
        plt.savefig(hist_path, dpi=100, bbox_inches='tight')
        plt.close()
        
        # Create boxplot
        plt.figure(figsize=(6, 2))
        sns.boxplot(x=data, color='lightgreen')
        plt.title(f"{metric_name} Boxplot")
        box_path = self.output_dir / f"{metric_name}_box.png"
        plt.savefig(box_path, dpi=100, bbox_inches='tight')
        plt.close()
        
        return hist_path, box_path
    
    def add_statistics_table(self, doc, metric_name, stats):
        """
        Add optimized statistics table to Word document.
        
        Shows 5 key metrics:
        - count: Total samples
        - mean: Average value
        - std: Standard deviation (consistency)
        - min: Minimum value (worst case)
        - max: Maximum value (best case)
        """
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        headers = ['Count', 'Mean', 'Std Dev', 'Min', 'Max']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        # Data row
        data_cells = table.rows[1].cells
        stat_keys = ['count', 'mean', 'std', 'min', 'max']
        for i, key in enumerate(stat_keys):
            value = stats.get(key, 0)
            # Format count as integer, others as decimal
            if key == 'count':
                data_cells[i].text = str(int(value))
            else:
                data_cells[i].text = f"{value:.4f}"
    
    def get_sample_images(self, n_samples=None):
        """Get sample image paths from database"""
        if n_samples is None:
            n_samples = self.n_sample_images
        if 'file_path' in self.df.columns:
            sample_paths = self.df['file_path'].dropna().head(n_samples)
            # Filter existing paths
            existing_paths = [path for path in sample_paths if os.path.exists(path)]
            return existing_paths
        return []
    
    def classify_quality(self, stats_summary):
        """
        Classify images into quality categories based on metrics.
        
        Quality classification criteria:
        - Good: High sharpness (laplacian, tenengrad), low noise
        - Fair: Medium values
        - Poor: Low sharpness, high noise
        
        Returns:
            dict: Quality distribution with counts and percentages
        """
        # Define thresholds based on typical endoscopic image characteristics
        quality_categories = {'Good': 0, 'Fair': 0, 'Poor': 0}
        
        if 'laplacian_variance' not in self.df.columns or 'noise_estimate' not in self.df.columns:
            return quality_categories
        
        for _, row in self.df.iterrows():
            laplacian = row.get('laplacian_variance', 0)
            noise = row.get('noise_estimate', float('inf'))
            
            # Use mean and std for classification thresholds
            laplacian_mean = stats_summary.get('laplacian', {}).get('mean', 0)
            laplacian_std = stats_summary.get('laplacian', {}).get('std', 1)
            noise_mean = stats_summary.get('noise', {}).get('mean', 0)
            noise_std = stats_summary.get('noise', {}).get('std', 1)
            
            # Good quality: high sharpness, low noise
            if laplacian > (laplacian_mean + 0.5 * laplacian_std) and noise < (noise_mean - 0.3 * noise_std):
                quality_categories['Good'] += 1
            # Poor quality: low sharpness or high noise
            elif laplacian < (laplacian_mean - 0.5 * laplacian_std) or noise > (noise_mean + 0.5 * noise_std):
                quality_categories['Poor'] += 1
            else:
                quality_categories['Fair'] += 1
        
        # Calculate percentages
        total = sum(quality_categories.values())
        if total > 0:
            quality_percentages = {
                k: {'count': v, 'percentage': (v / total) * 100} 
                for k, v in quality_categories.items()
            }
            return quality_percentages
        
        return quality_categories
    
    def calculate_correlation_matrix(self):
        """
        Calculate correlation matrix between metrics.
        
        Useful for understanding metric relationships and redundancy.
        
        Returns:
            pandas.DataFrame: Correlation matrix
        """
        # Select only numeric columns for metrics
        metric_cols = [col for col in self.metrics.keys() if col in self.df.columns]
        
        if len(metric_cols) < 2:
            return None
        
        correlation_matrix = self.df[metric_cols].corr()
        
        # Rename columns to metric names
        correlation_matrix.columns = [self.metrics.get(col, col) for col in correlation_matrix.columns]
        correlation_matrix.index = [self.metrics.get(col, col) for col in correlation_matrix.index]
        
        return correlation_matrix
    
    def create_correlation_heatmap(self, correlation_matrix):
        """
        Create correlation heatmap visualization.
        
        Args:
            correlation_matrix: Correlation matrix DataFrame
            
        Returns:
            Path to saved heatmap image
        """
        if correlation_matrix is None or correlation_matrix.empty:
            return None
        
        plt.figure(figsize=(10, 8))
        sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm', center=0, 
                   fmt='.2f', square=True, linewidths=0.5)
        plt.title("Metric Correlation Matrix", fontsize=14, fontweight='bold')
        plt.tight_layout()
        
        heatmap_path = self.output_dir / "correlation_heatmap.png"
        plt.savefig(heatmap_path, dpi=150, bbox_inches='tight')
        plt.close()
        
        return heatmap_path
    
    def add_summary_dashboard(self, doc, stats_summary, quality_distribution):
        """
        Add executive summary dashboard section.
        
        This provides a high-level overview with key insights:
        - Dataset overview
        - Quality distribution
        - Key findings
        """
        doc.add_heading("📊 Executive Summary Dashboard", level=1)
        
        # Dataset overview
        doc.add_heading("Dataset Overview", level=2)
        overview_table = doc.add_table(rows=2, cols=2)
        overview_table.style = 'Light List Accent 1'
        
        overview_table.rows[0].cells[0].text = "Total Images"
        overview_table.rows[0].cells[1].text = str(len(self.df))
        overview_table.rows[1].cells[0].text = "Metrics Analyzed"
        overview_table.rows[1].cells[1].text = str(len(stats_summary))
        
        # Quality distribution
        doc.add_heading("Quality Distribution", level=2)
        if isinstance(quality_distribution, dict) and any('count' in v for v in quality_distribution.values() if isinstance(v, dict)):
            quality_table = doc.add_table(rows=4, cols=3)
            quality_table.style = 'Light Grid Accent 1'
            
            # Headers
            quality_table.rows[0].cells[0].text = "Category"
            quality_table.rows[0].cells[1].text = "Count"
            quality_table.rows[0].cells[2].text = "Percentage"
            
            # Data rows
            categories = ['Good', 'Fair', 'Poor']
            for i, category in enumerate(categories, start=1):
                data = quality_distribution.get(category, {'count': 0, 'percentage': 0})
                quality_table.rows[i].cells[0].text = category
                quality_table.rows[i].cells[1].text = str(data.get('count', 0))
                quality_table.rows[i].cells[2].text = f"{data.get('percentage', 0):.1f}%"
        
        # Key findings
        doc.add_heading("Key Findings", level=2)
        findings = []
        
        # Analyze sharpness metrics
        if 'laplacian' in stats_summary:
            laplacian_mean = stats_summary['laplacian']['mean']
            findings.append(f"Average sharpness (Laplacian): {laplacian_mean:.4f}")
        
        # Analyze noise
        if 'noise' in stats_summary:
            noise_mean = stats_summary['noise']['mean']
            findings.append(f"Average noise level: {noise_mean:.4f}")
        
        # Analyze contrast
        if 'rms_contrast' in stats_summary:
            contrast_mean = stats_summary['rms_contrast']['mean']
            findings.append(f"Average contrast (RMS): {contrast_mean:.4f}")
        
        # Add findings to document
        for finding in findings:
            doc.add_paragraph(f"• {finding}", style='List Bullet')
        
        doc.add_page_break()
    
    def generate_report(self):
        """
        Generate optimized Word report with dashboard-ready metrics.
        
        Report structure:
        1. Executive Summary Dashboard (overview, quality distribution, key findings)
        2. Correlation Analysis (relationships between metrics)
        3. Individual Metric Details (5 key statistics + visualizations)
        """
        print("\n" + "="*60)
        print("GENERATING OPTIMIZED STATISTICAL REPORT")
        print("="*60)
        
        # Connect and fetch data
        self.connect_database()
        self.fetch_data()
        
        if self.df is None or len(self.df) == 0:
            print("✗ No data available to generate report")
            return
        
        # Calculate statistics
        print("\n📊 Calculating statistics...")
        stats_summary = self.calculate_statistics()
        
        # Classify quality
        print("📈 Classifying image quality...")
        quality_distribution = self.classify_quality(stats_summary)
        
        # Calculate correlations
        print("🔗 Calculating metric correlations...")
        correlation_matrix = self.calculate_correlation_matrix()
        
        # Create Word document
        print("\n📄 Creating optimized Word document...")
        doc = Document()
        doc.add_heading("Endoskopik Tasvir Sifat Statistikasi", 0)
        doc.add_paragraph(f"Hisobot yaratilgan sana: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        
        # Add Executive Summary Dashboard
        print("  Adding executive summary dashboard...")
        self.add_summary_dashboard(doc, stats_summary, quality_distribution)
        
        # Add Correlation Analysis
        if correlation_matrix is not None:
            print("  Adding correlation analysis...")
            doc.add_heading("🔗 Metric Correlation Analysis", level=1)
            doc.add_paragraph(
                "Korrelyatsiya tahlili metrikalar orasidagi bog'liqlikni ko'rsatadi. "
                "Yuqori korrelyatsiya (>0.7) metrikalar o'xshash axborotni beradi."
            )
            
            # Create and add heatmap
            heatmap_path = self.create_correlation_heatmap(correlation_matrix)
            if heatmap_path and os.path.exists(heatmap_path):
                doc.add_picture(str(heatmap_path), width=Inches(6))
            
            doc.add_page_break()
        
        # Generate sections for each metric
        doc.add_heading("📋 Detailed Metric Analysis", level=1)
        doc.add_paragraph(
            "Har bir metrika uchun 5 ta asosiy statistik ko'rsatkich:\n"
            "• Count: Namunalar soni\n"
            "• Mean: O'rtacha qiymat\n"
            "• Std Dev: Standart og'ish (izchillik ko'rsatkichi)\n"
            "• Min: Minimal qiymat (eng yomon holat)\n"
            "• Max: Maksimal qiymat (eng yaxshi holat)\n"
        )
        doc.add_page_break()
        
        for db_col, metric_name in self.metrics.items():
            if metric_name not in stats_summary:
                continue
            
            print(f"  Processing: {metric_name}")
            
            # Add metric heading and description
            doc.add_heading(metric_name.upper(), level=1)
            doc.add_paragraph(self.metric_descriptions.get(metric_name, ""))
            
            # Add statistics table
            self.add_statistics_table(doc, metric_name, stats_summary[metric_name])
            
            # Generate and add visualizations
            hist_path, box_path = self.create_visualizations(metric_name, db_col)
            
            if hist_path and os.path.exists(hist_path):
                doc.add_paragraph("\nTaqsimot grafigi:")
                doc.add_picture(str(hist_path), width=Inches(5))
            
            if box_path and os.path.exists(box_path):
                doc.add_paragraph("\nBoxplot (dispersiya va outlierlar):")
                doc.add_picture(str(box_path), width=Inches(5))
            
            doc.add_page_break()
        
        # Save document
        output_path = self.output_dir / "Image_Quality_Report_Optimized.docx"
        doc.save(str(output_path))
        
        print("\n" + "="*60)
        print(f"✓ Optimized report saved to: {output_path}")
        print(f"✓ Visualizations saved to: {self.output_dir}")
        print(f"✓ Report includes:")
        print(f"  - Executive summary dashboard")
        print(f"  - Quality classification")
        print(f"  - Correlation analysis")
        print(f"  - 5 key statistics per metric (reduced from 9)")
        print(f"  - Detailed visualizations")
        print("="*60)


def main():
    """
    Main function to generate optimized statistical report.
    
    Improvements:
    1. Reduced statistics from 9 to 5 most useful metrics
    2. Added executive summary dashboard
    3. Added quality classification (Good/Fair/Poor)
    4. Added correlation analysis between metrics
    5. Better organized for dashboard visualization
    
    Configuration options:
    - output_dir: Directory to save report and visualizations
    - n_sample_images: Number of sample images to include (default: 3)
    - histogram_bins: Number of bins for histograms or 'auto' (default: 'auto')
    """
    generator = StatisticsReportGenerator(
        DB_CONFIG, 
        output_dir='reports',
        n_sample_images=3,
        histogram_bins='auto'
    )
    
    # Generate optimized report
    generator.generate_report()
    
    print("\n✓ Optimized Word report ready!")
    print("\n📊 Report Features:")
    print("  • 5 key statistics (count, mean, std, min, max)")
    print("  • Executive summary dashboard")
    print("  • Quality distribution analysis")
    print("  • Metric correlation matrix")
    print("  • Detailed visualizations per metric")


if __name__ == "__main__":
    main()
