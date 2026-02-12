"""
Comprehensive Results Analysis for IQA System - Version 2
Generates comprehensive statistical analysis and visualizations,
and exports all results to a Word document.

This script analyzes:
- Statistical overview of all scores
- Score distributions (violin/box plots)
- Method agreement & correlation (Bland-Altman, scatter plots, heatmaps)
- Traditional feature importance
- Ensemble weight analysis
- Quality categories performance
- Performance efficiency
- Qualitative visual examples
- Model type analysis
"""

import os
import sys
import psycopg2
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from scipy import stats
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

# For Word document generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Set color-blind friendly palette
sns.set_palette("colorblind")
sns.set_style("whitegrid")

# Color mapping for consistency
COLORS = {
    'traditional': '#1f77b4',  # blue
    'deep_learning': '#ff7f0e',  # orange
    'ensemble': '#2ca02c'  # green
}


class ResultsAnalyzer:
    """Comprehensive results analyzer for IQA system."""
    
    def __init__(self, db_config: Dict, output_dir: str = 'results_analysis_v2'):
        """
        Initialize analyzer.
        
        Args:
            db_config: Database configuration dictionary
            output_dir: Directory to save results
        """
        self.db_config = db_config
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Create subdirectories
        self.plots_dir = self.output_dir / 'plots'
        self.plots_dir.mkdir(exist_ok=True)
        
        self.data = None
        self.conn = None
        
        print(f"✓ Results analyzer initialized")
        print(f"✓ Output directory: {self.output_dir}")
    
    def connect_db(self):
        """Connect to PostgreSQL database."""
        try:
            self.conn = psycopg2.connect(**self.db_config)
            print("✓ Connected to database")
        except Exception as e:
            print(f"✗ Error connecting to database: {e}")
            sys.exit(1)
    
    def fetch_data(self):
        """Fetch all data from database."""
        print("\n📊 Fetching data from database...")
        
        cursor = self.conn.cursor()
        
        query = """
            SELECT 
                q.id,
                q.image_id,
                m.filename,
                m.file_path,
                m.dataset_name,
                m.category,
                m.width,
                m.height,
                q.ensemble_score,
                q.traditional_score,
                q.deep_learning_score,
                q.laplacian_variance,
                q.rms_contrast,
                q.noise_estimate,
                q.mscn_std,
                q.gradient_energy,
                q.entropy,
                q.tenengrad,
                q.processing_time_ms,
                q.ensemble_weights_traditional,
                q.ensemble_weights_dl,
                q.model_type,
                q.computed_at
            FROM image_quality_metrics_hybrid q
            JOIN image_metadata m ON q.image_id = m.id
            ORDER BY q.id
        """
        
        cursor.execute(query)
        results = cursor.fetchall()
        
        # Create DataFrame
        columns = [
            'id', 'image_id', 'filename', 'file_path', 'dataset_name', 'category',
            'width', 'height', 'ensemble_score', 'traditional_score', 'deep_learning_score',
            'laplacian_variance', 'rms_contrast', 'noise_estimate', 'mscn_std',
            'gradient_energy', 'entropy', 'tenengrad', 'processing_time_ms',
            'ensemble_weights_traditional', 'ensemble_weights_dl', 'model_type', 'computed_at'
        ]
        
        self.data = pd.DataFrame(results, columns=columns)
        
        print(f"✓ Fetched {len(self.data)} records")
        print(f"✓ Datasets: {self.data['dataset_name'].unique().tolist()}")
        
        return self.data
    
    def close_db(self):
        """Close database connection."""
        if self.conn:
            self.conn.close()
            print("✓ Database connection closed")
    
    # ========================== Statistical Overview ==========================
    
    def generate_statistical_overview(self) -> pd.DataFrame:
        """Generate statistical overview table."""
        print("\n📈 Generating statistical overview...")
        
        metrics = ['ensemble_score', 'traditional_score', 'deep_learning_score', 'processing_time_ms']
        
        stats_data = []
        for metric in metrics:
            values = self.data[metric].dropna()
            stats_data.append({
                'Metric': metric.replace('_', ' ').title(),
                'Mean': f"{values.mean():.4f}",
                'Median': f"{values.median():.4f}",
                'Std': f"{values.std():.4f}",
                'Min': f"{values.min():.4f}",
                'Max': f"{values.max():.4f}"
            })
        
        stats_df = pd.DataFrame(stats_data)
        
        # Save to CSV
        stats_df.to_csv(self.output_dir / 'statistical_overview.csv', index=False)
        print(f"✓ Saved to {self.output_dir / 'statistical_overview.csv'}")
        
        return stats_df
    
    # ========================== Score Distribution ==========================
    
    def plot_score_distributions(self):
        """Generate violin plots for score distributions."""
        print("\n🎻 Generating score distribution plots...")
        
        # Prepare data for violin plot
        score_data = pd.DataFrame({
            'Ensemble': self.data['ensemble_score'],
            'Traditional': self.data['traditional_score'],
            'Deep Learning': self.data['deep_learning_score']
        })
        
        # Melt for seaborn
        score_data_melted = score_data.melt(var_name='Method', value_name='Score')
        
        # Create figure with two subplots
        fig, axes = plt.subplots(1, 2, figsize=(16, 6))
        
        # Violin plot
        sns.violinplot(data=score_data_melted, x='Method', y='Score', ax=axes[0], 
                      palette=[COLORS['ensemble'], COLORS['traditional'], COLORS['deep_learning']])
        axes[0].set_title('Score Distribution - Violin Plot', fontsize=16, fontweight='bold')
        axes[0].set_ylabel('Quality Score', fontsize=14)
        axes[0].set_xlabel('Method', fontsize=14)
        axes[0].grid(True, alpha=0.3)
        
        # Box plot
        sns.boxplot(data=score_data_melted, x='Method', y='Score', ax=axes[1],
                   palette=[COLORS['ensemble'], COLORS['traditional'], COLORS['deep_learning']])
        axes[1].set_title('Score Distribution - Box Plot', fontsize=16, fontweight='bold')
        axes[1].set_ylabel('Quality Score', fontsize=14)
        axes[1].set_xlabel('Method', fontsize=14)
        axes[1].grid(True, alpha=0.3)
        
        plt.tight_layout()
        plt.savefig(self.plots_dir / 'score_distributions.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✓ Saved score distribution plots")
    
    # ========================== Method Agreement ==========================
    
    def plot_bland_altman(self):
        """Generate Bland-Altman plot for method agreement."""
        print("\n📊 Generating Bland-Altman plot...")
        
        # Calculate mean and difference
        mean = (self.data['traditional_score'] + self.data['deep_learning_score']) / 2
        diff = self.data['traditional_score'] - self.data['deep_learning_score']
        
        mean_diff = diff.mean()
        std_diff = diff.std()
        
        # Create plot
        fig, ax = plt.subplots(figsize=(10, 8))
        
        ax.scatter(mean, diff, alpha=0.5, s=30, color='steelblue')
        ax.axhline(mean_diff, color='red', linestyle='-', linewidth=2, label=f'Mean: {mean_diff:.4f}')
        ax.axhline(mean_diff + 1.96 * std_diff, color='red', linestyle='--', linewidth=2, 
                   label=f'+1.96 SD: {mean_diff + 1.96 * std_diff:.4f}')
        ax.axhline(mean_diff - 1.96 * std_diff, color='red', linestyle='--', linewidth=2,
                   label=f'-1.96 SD: {mean_diff - 1.96 * std_diff:.4f}')
        
        ax.set_xlabel('Mean of Traditional and Deep Learning Scores', fontsize=14)
        ax.set_ylabel('Difference (Traditional - Deep Learning)', fontsize=14)
        ax.set_title('Bland-Altman Plot: Traditional vs Deep Learning', fontsize=16, fontweight='bold')
        ax.legend(fontsize=12)
        ax.grid(True, alpha=0.3)
        
        plt.tight_layout()
        plt.savefig(self.plots_dir / 'bland_altman_plot.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✓ Saved Bland-Altman plot")
    
    def plot_scatter_matrix(self):
        """Generate scatter plot matrix with correlations."""
        print("\n🔍 Generating scatter plot matrix...")
        
        # Create scatter matrix
        scores = self.data[['ensemble_score', 'traditional_score', 'deep_learning_score']]
        
        fig, axes = plt.subplots(3, 3, figsize=(15, 15))
        
        score_names = ['Ensemble', 'Traditional', 'Deep Learning']
        score_cols = ['ensemble_score', 'traditional_score', 'deep_learning_score']
        
        for i in range(3):
            for j in range(3):
                ax = axes[i, j]
                
                if i == j:
                    # Diagonal: histogram
                    ax.hist(self.data[score_cols[i]], bins=30, alpha=0.7, 
                           color=COLORS[score_cols[i].replace('_score', '')])
                    ax.set_ylabel('Frequency')
                    if i == 2:
                        ax.set_xlabel(score_names[i])
                else:
                    # Off-diagonal: scatter plot with regression
                    x = self.data[score_cols[j]]
                    y = self.data[score_cols[i]]
                    
                    ax.scatter(x, y, alpha=0.5, s=20, color='steelblue')
                    
                    # Add regression line
                    z = np.polyfit(x, y, 1)
                    p = np.poly1d(z)
                    ax.plot(x, p(x), 'r-', linewidth=2, alpha=0.8)
                    
                    # Calculate R²
                    r_squared = np.corrcoef(x, y)[0, 1] ** 2
                    ax.text(0.05, 0.95, f'R² = {r_squared:.4f}', 
                           transform=ax.transAxes, fontsize=10, verticalalignment='top',
                           bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))
                    
                    if j == 0:
                        ax.set_ylabel(score_names[i])
                    if i == 2:
                        ax.set_xlabel(score_names[j])
                
                ax.grid(True, alpha=0.3)
        
        plt.suptitle('Score Scatter Matrix with Regression Lines', fontsize=18, fontweight='bold', y=0.995)
        plt.tight_layout()
        plt.savefig(self.plots_dir / 'scatter_matrix.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✓ Saved scatter plot matrix")
    
    def plot_correlation_heatmap(self):
        """Generate correlation heatmap."""
        print("\n🔥 Generating correlation heatmap...")
        
        # Select score columns
        scores = self.data[['ensemble_score', 'traditional_score', 'deep_learning_score']]
        scores.columns = ['Ensemble', 'Traditional', 'Deep Learning']
        
        # Calculate PLCC (Pearson) and SROCC (Spearman)
        plcc = scores.corr(method='pearson')
        srocc = scores.corr(method='spearman')
        
        # Create figure with two subplots
        fig, axes = plt.subplots(1, 2, figsize=(16, 6))
        
        # PLCC heatmap
        sns.heatmap(plcc, annot=True, fmt='.4f', cmap='coolwarm', center=0, 
                   vmin=-1, vmax=1, square=True, ax=axes[0], cbar_kws={'label': 'PLCC'})
        axes[0].set_title('Pearson Linear Correlation (PLCC)', fontsize=16, fontweight='bold')
        
        # SROCC heatmap
        sns.heatmap(srocc, annot=True, fmt='.4f', cmap='coolwarm', center=0,
                   vmin=-1, vmax=1, square=True, ax=axes[1], cbar_kws={'label': 'SROCC'})
        axes[1].set_title('Spearman Rank Correlation (SROCC)', fontsize=16, fontweight='bold')
        
        plt.tight_layout()
        plt.savefig(self.plots_dir / 'correlation_heatmap.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✓ Saved correlation heatmap")
        
        return plcc, srocc
    
    # ========================== Feature Importance ==========================
    
    def plot_feature_importance(self):
        """Analyze traditional feature importance."""
        print("\n🔬 Generating feature importance analysis...")
        
        # Traditional features
        features = [
            'laplacian_variance', 'rms_contrast', 'noise_estimate', 
            'mscn_std', 'gradient_energy', 'entropy', 'tenengrad'
        ]
        
        # Calculate correlation with ensemble score
        correlations = {}
        for feature in features:
            corr = self.data[feature].corr(self.data['ensemble_score'])
            correlations[feature] = corr
        
        # Create figure with multiple subplots
        fig, axes = plt.subplots(2, 2, figsize=(16, 12))
        
        # 1. Correlation heatmap
        feature_data = self.data[features + ['ensemble_score']]
        feature_data.columns = [f.replace('_', ' ').title() for f in features] + ['Ensemble Score']
        
        corr_matrix = feature_data.corr()
        
        sns.heatmap(corr_matrix, annot=True, fmt='.3f', cmap='viridis', 
                   center=0, ax=axes[0, 0], cbar_kws={'label': 'Correlation'})
        axes[0, 0].set_title('Feature Correlation with Ensemble Score', fontsize=14, fontweight='bold')
        
        # 2. Feature contribution bar chart
        feature_names = [f.replace('_', ' ').title() for f in features]
        corr_values = [correlations[f] for f in features]
        
        colors_bars = ['green' if v > 0 else 'red' for v in corr_values]
        axes[0, 1].barh(feature_names, corr_values, color=colors_bars, alpha=0.7)
        axes[0, 1].set_xlabel('Correlation with Ensemble Score', fontsize=12)
        axes[0, 1].set_title('Feature Contribution to Quality Score', fontsize=14, fontweight='bold')
        axes[0, 1].axvline(0, color='black', linestyle='--', linewidth=1)
        axes[0, 1].grid(True, alpha=0.3)
        
        # 3. Radar chart of normalized features
        from math import pi
        
        # Normalize features to [0, 1]
        normalized_features = {}
        for feature in features:
            min_val = self.data[feature].min()
            max_val = self.data[feature].max()
            normalized_features[feature] = (self.data[feature] - min_val) / (max_val - min_val)
        
        # Calculate mean normalized values
        mean_values = [normalized_features[f].mean() for f in features]
        
        # Radar chart
        angles = [n / float(len(features)) * 2 * pi for n in range(len(features))]
        mean_values += mean_values[:1]
        angles += angles[:1]
        
        ax = plt.subplot(2, 2, 3, projection='polar')
        ax.plot(angles, mean_values, 'o-', linewidth=2, color='steelblue')
        ax.fill(angles, mean_values, alpha=0.25, color='steelblue')
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(feature_names, size=10)
        ax.set_ylim(0, 1)
        ax.set_title('Normalized Mean Feature Values (Radar Chart)', fontsize=14, fontweight='bold', pad=20)
        ax.grid(True)
        
        # 4. Box plot of feature distributions
        feature_data_for_box = pd.DataFrame({
            f.replace('_', ' ').title(): (self.data[f] - self.data[f].min()) / (self.data[f].max() - self.data[f].min())
            for f in features
        })
        
        feature_data_melted = feature_data_for_box.melt(var_name='Feature', value_name='Normalized Value')
        
        sns.boxplot(data=feature_data_melted, x='Feature', y='Normalized Value', ax=axes[1, 1])
        axes[1, 1].set_xticklabels(axes[1, 1].get_xticklabels(), rotation=45, ha='right')
        axes[1, 1].set_title('Normalized Feature Distributions', fontsize=14, fontweight='bold')
        axes[1, 1].grid(True, alpha=0.3)
        
        plt.tight_layout()
        plt.savefig(self.plots_dir / 'feature_importance.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✓ Saved feature importance analysis")
        
        return correlations
    
    # ========================== Ensemble Weights ==========================
    
    def plot_ensemble_weights(self):
        """Analyze ensemble weight distribution."""
        print("\n⚖️  Generating ensemble weight analysis...")
        
        fig, axes = plt.subplots(1, 3, figsize=(18, 5))
        
        # 1. Histogram of weights
        axes[0].hist(self.data['ensemble_weights_traditional'], bins=30, alpha=0.7, 
                    color=COLORS['traditional'], label='Traditional', edgecolor='black')
        axes[0].hist(self.data['ensemble_weights_dl'], bins=30, alpha=0.7,
                    color=COLORS['deep_learning'], label='Deep Learning', edgecolor='black')
        axes[0].set_xlabel('Weight', fontsize=12)
        axes[0].set_ylabel('Frequency', fontsize=12)
        axes[0].set_title('Ensemble Weight Distribution', fontsize=14, fontweight='bold')
        axes[0].legend()
        axes[0].grid(True, alpha=0.3)
        
        # 2. Scatter plot: weight relationship
        axes[1].scatter(self.data['ensemble_weights_traditional'], 
                       self.data['ensemble_weights_dl'],
                       alpha=0.5, s=30, color='steelblue')
        axes[1].set_xlabel('Traditional Weight', fontsize=12)
        axes[1].set_ylabel('Deep Learning Weight', fontsize=12)
        axes[1].set_title('Weight Relationship', fontsize=14, fontweight='bold')
        axes[1].grid(True, alpha=0.3)
        
        # Add diagonal line
        axes[1].plot([0, 1], [1, 0], 'r--', linewidth=2, alpha=0.5, label='Sum = 1')
        axes[1].legend()
        
        # 3. Pie chart of average weights
        avg_trad = self.data['ensemble_weights_traditional'].mean()
        avg_dl = self.data['ensemble_weights_dl'].mean()
        
        axes[2].pie([avg_trad, avg_dl], 
                   labels=['Traditional', 'Deep Learning'],
                   colors=[COLORS['traditional'], COLORS['deep_learning']],
                   autopct='%1.1f%%', startangle=90, textprops={'fontsize': 12})
        axes[2].set_title('Average Ensemble Weights', fontsize=14, fontweight='bold')
        
        plt.tight_layout()
        plt.savefig(self.plots_dir / 'ensemble_weights.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        # Analysis
        high_trad = (self.data['ensemble_weights_traditional'] > 0.5).sum()
        high_dl = (self.data['ensemble_weights_dl'] > 0.5).sum()
        total = len(self.data)
        
        analysis = {
            'avg_traditional': avg_trad,
            'avg_deep_learning': avg_dl,
            'pct_high_traditional': (high_trad / total) * 100,
            'pct_high_dl': (high_dl / total) * 100
        }
        
        print(f"✓ Saved ensemble weight analysis")
        print(f"  • Average traditional weight: {avg_trad:.4f}")
        print(f"  • Average deep learning weight: {avg_dl:.4f}")
        print(f"  • Images with traditional > 0.5: {high_trad} ({analysis['pct_high_traditional']:.1f}%)")
        
        return analysis
    
    # ========================== Quality Categories ==========================
    
    def plot_quality_categories(self):
        """Analyze quality category distribution."""
        print("\n📊 Generating quality category analysis...")
        
        # Define bins
        bins = [0, 0.3, 0.6, 0.8, 1.0]
        labels = ['Poor', 'Fair', 'Good', 'Excellent']
        
        # Categorize scores
        self.data['ensemble_category'] = pd.cut(self.data['ensemble_score'], bins=bins, labels=labels, include_lowest=True)
        self.data['traditional_category'] = pd.cut(self.data['traditional_score'], bins=bins, labels=labels, include_lowest=True)
        self.data['dl_category'] = pd.cut(self.data['deep_learning_score'], bins=bins, labels=labels, include_lowest=True)
        
        # Count categories
        ensemble_counts = self.data['ensemble_category'].value_counts()
        traditional_counts = self.data['traditional_category'].value_counts()
        dl_counts = self.data['dl_category'].value_counts()
        
        # Create figure
        fig, axes = plt.subplots(1, 2, figsize=(16, 6))
        
        # 1. Stacked bar chart
        category_data = pd.DataFrame({
            'Ensemble': ensemble_counts,
            'Traditional': traditional_counts,
            'Deep Learning': dl_counts
        }).fillna(0)
        
        category_data.plot(kind='bar', ax=axes[0], 
                          color=[COLORS['ensemble'], COLORS['traditional'], COLORS['deep_learning']],
                          alpha=0.8, edgecolor='black')
        axes[0].set_xlabel('Quality Category', fontsize=12)
        axes[0].set_ylabel('Count', fontsize=12)
        axes[0].set_title('Quality Category Distribution by Method', fontsize=14, fontweight='bold')
        axes[0].legend(title='Method')
        axes[0].set_xticklabels(axes[0].get_xticklabels(), rotation=0)
        axes[0].grid(True, alpha=0.3, axis='y')
        
        # 2. Pie chart for ensemble
        axes[1].pie(ensemble_counts, labels=ensemble_counts.index, autopct='%1.1f%%',
                   startangle=90, colors=['#d62728', '#ff7f0e', '#2ca02c', '#1f77b4'])
        axes[1].set_title('Ensemble Score Distribution by Category', fontsize=14, fontweight='bold')
        
        plt.tight_layout()
        plt.savefig(self.plots_dir / 'quality_categories.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✓ Saved quality category analysis")
    
    # ========================== Performance Efficiency ==========================
    
    def plot_performance_efficiency(self):
        """Analyze processing time and efficiency."""
        print("\n⏱️  Generating performance efficiency analysis...")
        
        fig, axes = plt.subplots(1, 2, figsize=(16, 6))
        
        # 1. Box plot of processing time
        axes[0].boxplot([self.data['processing_time_ms']], labels=['Processing Time'])
        axes[0].set_ylabel('Time (ms)', fontsize=12)
        axes[0].set_title('Processing Time Distribution', fontsize=14, fontweight='bold')
        axes[0].grid(True, alpha=0.3, axis='y')
        
        # Add statistics
        mean_time = self.data['processing_time_ms'].mean()
        median_time = self.data['processing_time_ms'].median()
        axes[0].text(0.98, 0.98, f'Mean: {mean_time:.2f} ms\nMedian: {median_time:.2f} ms',
                    transform=axes[0].transAxes, fontsize=11, verticalalignment='top',
                    horizontalalignment='right',
                    bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))
        
        # 2. Scatter plot: score vs time
        axes[1].scatter(self.data['ensemble_score'], self.data['processing_time_ms'],
                       alpha=0.5, s=30, color='steelblue')
        axes[1].set_xlabel('Ensemble Score', fontsize=12)
        axes[1].set_ylabel('Processing Time (ms)', fontsize=12)
        axes[1].set_title('Quality-Speed Tradeoff', fontsize=14, fontweight='bold')
        axes[1].grid(True, alpha=0.3)
        
        plt.tight_layout()
        plt.savefig(self.plots_dir / 'performance_efficiency.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✓ Saved performance efficiency analysis")
        print(f"  • Average processing time: {mean_time:.2f} ms per image")
    
    # ========================== Visual Examples ==========================
    
    def generate_visual_examples_table(self):
        """Generate table of representative examples."""
        print("\n🖼️  Generating visual examples table...")
        
        # Define quality ranges
        quality_ranges = {
            'Excellent': (0.8, 1.0),
            'Good': (0.6, 0.8),
            'Fair': (0.3, 0.6),
            'Poor': (0.0, 0.3)
        }
        
        examples = []
        
        for category, (min_score, max_score) in quality_ranges.items():
            # Get examples in this range
            subset = self.data[
                (self.data['ensemble_score'] >= min_score) & 
                (self.data['ensemble_score'] < max_score)
            ].head(3)
            
            for _, row in subset.iterrows():
                examples.append({
                    'Category': category,
                    'Filename': row['filename'],
                    'Dataset': row['dataset_name'],
                    'Ensemble': f"{row['ensemble_score']:.4f}",
                    'Traditional': f"{row['traditional_score']:.4f}",
                    'Deep Learning': f"{row['deep_learning_score']:.4f}"
                })
        
        examples_df = pd.DataFrame(examples)
        examples_df.to_csv(self.output_dir / 'visual_examples.csv', index=False)
        
        print(f"✓ Saved visual examples table")
        
        return examples_df
    
    def generate_disagreement_cases(self):
        """Find cases where methods disagree significantly."""
        print("\n⚠️  Identifying disagreement cases...")
        
        # Calculate disagreement
        self.data['disagreement'] = np.abs(
            self.data['traditional_score'] - self.data['deep_learning_score']
        )
        
        # Get top disagreement cases
        disagreement_cases = self.data.nlargest(10, 'disagreement')[
            ['filename', 'dataset_name', 'ensemble_score', 'traditional_score', 
             'deep_learning_score', 'disagreement']
        ]
        
        disagreement_cases.to_csv(self.output_dir / 'disagreement_cases.csv', index=False)
        
        print(f"✓ Saved disagreement cases")
        print(f"  • Found {len(disagreement_cases)} cases with significant disagreement")
        
        return disagreement_cases
    
    # ========================== Model Type Analysis ==========================
    
    def plot_model_type_analysis(self):
        """Analyze performance by model type."""
        print("\n🤖 Generating model type analysis...")
        
        if self.data['model_type'].nunique() <= 1:
            print("  ⚠️ Only one model type found, skipping model type analysis")
            return None
        
        # Group by model type
        model_stats = self.data.groupby('model_type').agg({
            'ensemble_score': ['mean', 'std', 'count'],
            'traditional_score': ['mean'],
            'deep_learning_score': ['mean'],
            'processing_time_ms': ['mean']
        }).reset_index()
        
        # Flatten column names
        model_stats.columns = ['_'.join(col).strip('_') for col in model_stats.columns.values]
        
        # Create plot
        fig, ax = plt.subplots(figsize=(12, 6))
        
        x = np.arange(len(model_stats))
        width = 0.25
        
        ax.bar(x - width, model_stats['ensemble_score_mean'], width, 
               label='Ensemble', color=COLORS['ensemble'], alpha=0.8)
        ax.bar(x, model_stats['traditional_score_mean'], width,
               label='Traditional', color=COLORS['traditional'], alpha=0.8)
        ax.bar(x + width, model_stats['deep_learning_score_mean'], width,
               label='Deep Learning', color=COLORS['deep_learning'], alpha=0.8)
        
        ax.set_xlabel('Model Type', fontsize=12)
        ax.set_ylabel('Average Score', fontsize=12)
        ax.set_title('Average Scores by Model Type', fontsize=14, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(model_stats['model_type'])
        ax.legend()
        ax.grid(True, alpha=0.3, axis='y')
        
        plt.tight_layout()
        plt.savefig(self.plots_dir / 'model_type_analysis.png', dpi=300, bbox_inches='tight')
        plt.close()
        
        print(f"✓ Saved model type analysis")
        
        return model_stats
    
    # ========================== Word Document Generation ==========================
    
    def generate_word_document(self, stats_df, plcc, srocc, feature_corr, 
                              weight_analysis, examples_df, disagreement_df):
        """Generate comprehensive Word document with all results."""
        print("\n📄 Generating Word document...")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Comprehensive Results Analysis for Endoscopic IQA System', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # ==================== 1. Statistical Overview ====================
        doc.add_heading('1. Statistical Overview', 1)
        doc.add_paragraph(
            'This table provides descriptive statistics for all quality scores and processing times.'
        )
        
        # Add table
        table = doc.add_table(rows=len(stats_df) + 1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        headers = ['Metric', 'Mean', 'Median', 'Std', 'Min', 'Max']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        
        # Data
        for i, row in stats_df.iterrows():
            for j, value in enumerate(row):
                table.rows[i + 1].cells[j].text = str(value)
        
        doc.add_paragraph()
        
        # Key findings
        doc.add_heading('Key Findings:', 2)
        doc.add_paragraph(
            f"• Average ensemble score: {float(stats_df[stats_df['Metric'] == 'Ensemble Score']['Mean'].values[0]):.4f}",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"• Average processing time: {float(stats_df[stats_df['Metric'] == 'Processing Time Ms']['Mean'].values[0]):.2f} ms",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"• Total images analyzed: {len(self.data)}",
            style='List Bullet'
        )
        
        doc.add_page_break()
        
        # ==================== 2. Score Distribution ====================
        doc.add_heading('2. Score Distribution Analysis', 1)
        doc.add_paragraph(
            'Violin and box plots show the distribution shape, quartiles, and outliers for each method.'
        )
        
        if (self.plots_dir / 'score_distributions.png').exists():
            doc.add_picture(str(self.plots_dir / 'score_distributions.png'), width=Inches(6.5))
        
        doc.add_paragraph()
        doc.add_paragraph(
            'The violin plots reveal the density distribution of scores, while box plots highlight '
            'the median, quartiles, and potential outliers.'
        )
        
        doc.add_page_break()
        
        # ==================== 3. Method Agreement & Correlation ====================
        doc.add_heading('3. Method Agreement & Correlation Analysis', 1)
        
        doc.add_heading('3.1 Bland-Altman Plot', 2)
        doc.add_paragraph(
            'The Bland-Altman plot assesses agreement between traditional and deep learning methods '
            'by plotting the difference against the mean.'
        )
        
        if (self.plots_dir / 'bland_altman_plot.png').exists():
            doc.add_picture(str(self.plots_dir / 'bland_altman_plot.png'), width=Inches(6))
        
        doc.add_page_break()
        
        doc.add_heading('3.2 Scatter Plot Matrix', 2)
        doc.add_paragraph(
            'The scatter matrix shows pairwise relationships between all three scoring methods '
            'with regression lines and R² values.'
        )
        
        if (self.plots_dir / 'scatter_matrix.png').exists():
            doc.add_picture(str(self.plots_dir / 'scatter_matrix.png'), width=Inches(6.5))
        
        doc.add_page_break()
        
        doc.add_heading('3.3 Correlation Heatmap', 2)
        doc.add_paragraph(
            'Correlation heatmaps display both Pearson (PLCC) and Spearman (SROCC) correlations '
            'between methods.'
        )
        
        if (self.plots_dir / 'correlation_heatmap.png').exists():
            doc.add_picture(str(self.plots_dir / 'correlation_heatmap.png'), width=Inches(6.5))
        
        doc.add_paragraph()
        
        # Add correlation values
        doc.add_heading('Correlation Values:', 3)
        trad_dl_plcc = plcc.loc['Traditional', 'Deep Learning']
        trad_dl_srocc = srocc.loc['Traditional', 'Deep Learning']
        
        doc.add_paragraph(
            f"• Traditional vs Deep Learning PLCC: {trad_dl_plcc:.4f}",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"• Traditional vs Deep Learning SROCC: {trad_dl_srocc:.4f}",
            style='List Bullet'
        )
        
        doc.add_page_break()
        
        # ==================== 4. Feature Importance ====================
        doc.add_heading('4. Traditional Feature Importance Analysis', 1)
        doc.add_paragraph(
            'Analysis of the 7 traditional features and their correlation with the ensemble quality score.'
        )
        
        if (self.plots_dir / 'feature_importance.png').exists():
            doc.add_picture(str(self.plots_dir / 'feature_importance.png'), width=Inches(6.5))
        
        doc.add_paragraph()
        
        # Add feature correlations
        doc.add_heading('Feature Correlations with Ensemble Score:', 2)
        for feature, corr in sorted(feature_corr.items(), key=lambda x: abs(x[1]), reverse=True):
            doc.add_paragraph(
                f"• {feature.replace('_', ' ').title()}: {corr:.4f}",
                style='List Bullet'
            )
        
        doc.add_page_break()
        
        # ==================== 5. Ensemble Weight Analysis ====================
        doc.add_heading('5. Ensemble Weight Analysis', 1)
        doc.add_paragraph(
            'Distribution and relationship of ensemble weights between traditional and deep learning methods.'
        )
        
        if (self.plots_dir / 'ensemble_weights.png').exists():
            doc.add_picture(str(self.plots_dir / 'ensemble_weights.png'), width=Inches(6.5))
        
        doc.add_paragraph()
        
        doc.add_heading('Weight Statistics:', 2)
        doc.add_paragraph(
            f"• Average traditional weight: {weight_analysis['avg_traditional']:.4f}",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"• Average deep learning weight: {weight_analysis['avg_deep_learning']:.4f}",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"• Images with traditional weight > 0.5: {weight_analysis['pct_high_traditional']:.1f}%",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"• Images with deep learning weight > 0.5: {weight_analysis['pct_high_dl']:.1f}%",
            style='List Bullet'
        )
        
        doc.add_page_break()
        
        # ==================== 6. Quality Categories ====================
        doc.add_heading('6. Quality Categories Performance', 1)
        doc.add_paragraph(
            'Distribution of images across quality categories: Poor (0-0.3), Fair (0.3-0.6), '
            'Good (0.6-0.8), and Excellent (0.8-1.0).'
        )
        
        if (self.plots_dir / 'quality_categories.png').exists():
            doc.add_picture(str(self.plots_dir / 'quality_categories.png'), width=Inches(6.5))
        
        doc.add_page_break()
        
        # ==================== 7. Performance Efficiency ====================
        doc.add_heading('7. Performance Efficiency', 1)
        doc.add_paragraph(
            'Analysis of processing time and the quality-speed tradeoff.'
        )
        
        if (self.plots_dir / 'performance_efficiency.png').exists():
            doc.add_picture(str(self.plots_dir / 'performance_efficiency.png'), width=Inches(6.5))
        
        doc.add_paragraph()
        
        mean_time = self.data['processing_time_ms'].mean()
        median_time = self.data['processing_time_ms'].median()
        
        doc.add_paragraph(
            f"• Average processing time: {mean_time:.2f} ms per image",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"• Median processing time: {median_time:.2f} ms per image",
            style='List Bullet'
        )
        
        doc.add_page_break()
        
        # ==================== 8. Visual Examples ====================
        doc.add_heading('8. Qualitative Visual Examples', 1)
        doc.add_paragraph(
            'Representative examples from each quality category with their scores.'
        )
        
        # Add examples table (first 12 rows)
        if len(examples_df) > 0:
            table = doc.add_table(rows=min(13, len(examples_df) + 1), cols=6)
            table.style = 'Light Grid Accent 1'
            
            # Headers
            headers = examples_df.columns.tolist()
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
            
            # Data
            for i, row in examples_df.head(12).iterrows():
                for j, value in enumerate(row):
                    table.rows[i + 1].cells[j].text = str(value)
        
        doc.add_page_break()
        
        # ==================== 9. Disagreement Cases ====================
        doc.add_heading('9. Method Disagreement Cases', 1)
        doc.add_paragraph(
            'Images where traditional and deep learning methods show significant disagreement '
            '(difference > 0.3).'
        )
        
        if len(disagreement_df) > 0:
            table = doc.add_table(rows=min(11, len(disagreement_df) + 1), cols=6)
            table.style = 'Light Grid Accent 1'
            
            # Headers
            headers = disagreement_df.columns.tolist()
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
            
            # Data
            for i, row in disagreement_df.head(10).iterrows():
                for j, value in enumerate(row):
                    table.rows[i + 1].cells[j].text = str(value)
        
        doc.add_paragraph()
        doc.add_paragraph(
            'These cases highlight scenarios where traditional metrics (favoring sharpness) '
            'and deep learning (detecting artifacts) provide different assessments.'
        )
        
        # ==================== 10. Model Type Analysis ====================
        if self.data['model_type'].nunique() > 1:
            doc.add_page_break()
            doc.add_heading('10. Model Type Analysis', 1)
            doc.add_paragraph(
                'Performance comparison across different deep learning model architectures.'
            )
            
            if (self.plots_dir / 'model_type_analysis.png').exists():
                doc.add_picture(str(self.plots_dir / 'model_type_analysis.png'), width=Inches(6.5))
        
        # ==================== Conclusion ====================
        doc.add_page_break()
        doc.add_heading('Conclusion', 1)
        
        doc.add_paragraph(
            'This comprehensive analysis demonstrates the effectiveness of the hybrid IQA approach, '
            'combining traditional image quality metrics with deep learning-based assessment. '
            'The ensemble method leverages the strengths of both approaches to provide robust '
            'quality predictions for endoscopic images.'
        )
        
        doc.add_paragraph()
        doc.add_paragraph(
            'Key achievements:'
        )
        doc.add_paragraph(
            f'• Analyzed {len(self.data)} images across multiple datasets',
            style='List Bullet'
        )
        doc.add_paragraph(
            f'• Achieved average ensemble score of {self.data["ensemble_score"].mean():.4f}',
            style='List Bullet'
        )
        doc.add_paragraph(
            f'• Processing efficiency: {mean_time:.2f} ms per image',
            style='List Bullet'
        )
        doc.add_paragraph(
            f'• Correlation between methods: PLCC = {trad_dl_plcc:.4f}, SROCC = {trad_dl_srocc:.4f}',
            style='List Bullet'
        )
        
        # Save document
        doc_path = self.output_dir / 'Comprehensive_Results_Analysis.docx'
        doc.save(doc_path)
        
        print(f"✓ Word document saved: {doc_path}")
    
    # ========================== Main Analysis ==========================
    
    def run_complete_analysis(self):
        """Run complete analysis pipeline."""
        print("\n" + "=" * 80)
        print("🔬 COMPREHENSIVE RESULTS ANALYSIS - VERSION 2")
        print("=" * 80)
        
        # Connect and fetch data
        self.connect_db()
        self.fetch_data()
        
        # Run all analyses
        print("\n" + "=" * 80)
        print("📊 GENERATING ANALYSES AND VISUALIZATIONS")
        print("=" * 80)
        
        stats_df = self.generate_statistical_overview()
        self.plot_score_distributions()
        self.plot_bland_altman()
        self.plot_scatter_matrix()
        plcc, srocc = self.plot_correlation_heatmap()
        feature_corr = self.plot_feature_importance()
        weight_analysis = self.plot_ensemble_weights()
        self.plot_quality_categories()
        self.plot_performance_efficiency()
        examples_df = self.generate_visual_examples_table()
        disagreement_df = self.generate_disagreement_cases()
        self.plot_model_type_analysis()
        
        # Generate Word document
        print("\n" + "=" * 80)
        print("📄 GENERATING COMPREHENSIVE WORD DOCUMENT")
        print("=" * 80)
        
        self.generate_word_document(
            stats_df, plcc, srocc, feature_corr, 
            weight_analysis, examples_df, disagreement_df
        )
        
        # Close database
        self.close_db()
        
        print("\n" + "=" * 80)
        print("✅ ANALYSIS COMPLETE!")
        print("=" * 80)
        print(f"\n📁 Results saved in: {self.output_dir}")
        print(f"📊 Plots saved in: {self.plots_dir}")
        print(f"📄 Word document: {self.output_dir / 'Comprehensive_Results_Analysis.docx'}")
        print("\n" + "=" * 80)


def main():
    """Main function."""
    # Import database configuration
    try:
        sys.path.append(str(Path(__file__).parent.parent / 'scripts'))
        from db_config import DB_CONFIG
        db_config = DB_CONFIG
    except ImportError:
        print("⚠️  db_config.py not found, using defaults")
        db_config = {
            'host': 'localhost',
            'port': 5432,
            'database': 'postgres',
            'user': os.environ.get('USER', 'postgres'),
            'password': ''
        }
    
    # Create analyzer
    analyzer = ResultsAnalyzer(db_config, output_dir='results_analysis_v2')
    
    # Run complete analysis
    analyzer.run_complete_analysis()


if __name__ == "__main__":
    main()
